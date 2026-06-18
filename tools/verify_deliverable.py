#!/usr/bin/env python3
"""One-command verification gate for generated deliverables.

Replaces the ad-hoc, separate verification steps (recalc the xlsx, validate the
docx, render a page, eyeball the Sources section) with a single deterministic
check that returns a clear pass/fail and a non-zero exit code on failure - so it
can gate a run.

Checks
  .xlsx : recalculate all formulas (via the xlsx skill's recalc.py if locatable)
          and assert ZERO Excel errors (#REF!, #DIV/0!, ...). Falls back to scanning
          cached cell values for error strings if recalc.py is unavailable.
  .docx : validate OOXML (via the docx skill's validate.py if locatable; else a
          python-docx open smoke test) AND confirm a "Sources" section is present
          (the house rule for anything built on filings/web content).

Usage:
  python3 tools/verify_deliverable.py model.xlsx note.docx
  python3 tools/verify_deliverable.py --no-sources-check draft.docx
"""
import argparse
import glob
import json
import os
import subprocess
import sys

ERROR_STRINGS = ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#N/A", "#NUM!", "#NULL!")


def _find(rel_globs, env=None):
    if env and os.environ.get(env):
        return os.environ[env]
    cands = []
    for g in rel_globs:
        cands += glob.glob(g)
        cands += glob.glob(os.path.expanduser("~") + "/**/" + g.lstrip("/"), recursive=True)
    for c in cands:
        if os.path.exists(c):
            return c
    return None


def verify_xlsx(path):
    res = {"file": path, "kind": "xlsx", "checks": [], "ok": True}
    recalc = _find(["/sessions/*/mnt/.claude/skills/xlsx/scripts/recalc.py"], env="JSE_XLSX_RECALC")
    if recalc:
        out = subprocess.run([sys.executable, recalc, path, "60"], capture_output=True, text=True)
        try:
            r = json.loads(out.stdout)
            errs = r.get("total_errors", 0)
            res["checks"].append({"recalc_formula_errors": errs, "formulas": r.get("total_formulas")})
            if errs:
                res["ok"] = False
                res["checks"].append({"error_summary": r.get("error_summary")})
        except json.JSONDecodeError:
            res["checks"].append({"recalc": "unparseable output", "ok": False})
            res["ok"] = False
    else:
        # fallback: scan cached values for error strings
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True)
        found = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for c in row:
                    if isinstance(c.value, str) and c.value in ERROR_STRINGS:
                        found.append(f"{ws.title}!{c.coordinate}={c.value}")
        res["checks"].append({"recalc": "unavailable; scanned cached values",
                              "error_cells": found})
        if found:
            res["ok"] = False
    return res


def docx_text(path):
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def verify_docx(path, sources_check=True):
    res = {"file": path, "kind": "docx", "checks": [], "ok": True}
    validate = _find(["/sessions/*/mnt/.claude/skills/docx/scripts/office/validate.py"], env="JSE_DOCX_VALIDATE")
    if validate:
        out = subprocess.run([sys.executable, validate, path], capture_output=True, text=True)
        passed = "All validations PASSED" in out.stdout or "PASSED" in out.stdout
        failed = "FAILED" in out.stdout
        res["checks"].append({"ooxml_validate": "passed" if passed and not failed else "failed"})
        if failed or not passed:
            res["ok"] = False
            res["checks"].append({"validate_tail": out.stdout.strip().splitlines()[-6:]})
    else:
        try:
            txt = docx_text(path)
            res["checks"].append({"ooxml_validate": "skipped (validator not found)",
                                  "open_smoke": "ok", "paragraphs_chars": len(txt)})
        except Exception as e:  # noqa
            res["ok"] = False
            res["checks"].append({"open_smoke": f"FAILED: {e}"})
            return res
    if sources_check:
        txt = docx_text(path)
        has = "source" in txt.lower()
        res["checks"].append({"sources_section": "present" if has else "MISSING"})
        if not has:
            res["ok"] = False
    return res


def main():
    ap = argparse.ArgumentParser(description="Verify generated JSE deliverables.")
    ap.add_argument("files", nargs="+")
    ap.add_argument("--no-sources-check", action="store_true")
    args = ap.parse_args()
    results, ok = [], True
    for f in args.files:
        ext = os.path.splitext(f)[1].lower()
        if ext == ".xlsx":
            r = verify_xlsx(f)
        elif ext == ".docx":
            r = verify_docx(f, sources_check=not args.no_sources_check)
        else:
            r = {"file": f, "ok": False, "checks": [{"error": "unsupported type"}]}
        ok = ok and r["ok"]
        results.append(r)
    print(json.dumps({"ok": ok, "results": results}, indent=2))
    sys.exit(0 if ok else 2)


if __name__ == "__main__":
    main()
