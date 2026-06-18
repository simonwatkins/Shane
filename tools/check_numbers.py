#!/usr/bin/env python3
"""'Numbers are sacred' linter.

Operationalises core principle #1 of CLAUDE.md: every financial figure must either
be a reported actual taken from a cited source, or be explicitly marked (e) as an
analyst calculation/estimate. This tool scans a generated deliverable, extracts the
financial-looking numbers, and flags any that are NEITHER a known reported actual
(from the company's assumptions anchor / company.json key facts) NOR tagged as an
estimate (an "(e)" or "~" nearby). The flags are a review to-do, not a hard error.

It is intentionally conservative: it cannot know every legitimate number, so it
reports candidates for a human to confirm rather than asserting wrongness.

Usage:
  python3 tools/check_numbers.py note.docx --companies shoprite spar
  python3 tools/check_numbers.py model.xlsx --companies shoprite --strict
"""
import argparse
import json
import os
import re
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# financial-looking tokens: 1,234 / 1234.5 / 12.3% / R12.3bn / 795.8c / 1.74x
NUM_RE = re.compile(r"(?<![\w.])(?:R|US\$|\$|EUR|€)?\s?-?\d[\d,]*(?:\.\d+)?\s?(?:%|bn|m|c|x|bps|pp)?", re.I)


def normalise(tok):
    """Reduce a token to its bare numeric core for comparison."""
    t = tok.lower().replace(",", "").replace(" ", "")
    t = re.sub(r"^(r|us\$|\$|eur|€)", "", t)
    t = re.sub(r"(%|bn|m|c|x|bps|pp)$", "", t)
    t = t.lstrip("-")
    return t


def known_numbers(slugs):
    known = set()
    def add(v):
        if isinstance(v, (int, float)):
            for s in (f"{v}", f"{v:.1f}", f"{v:.2f}", f"{int(v)}" if float(v).is_integer() else f"{v}"):
                known.add(normalise(s))
            # percentages stored as fractions also appear x100
            known.add(normalise(f"{v*100:.1f}"))
            known.add(normalise(f"{v*100:.2f}"))
    for slug in slugs:
        ap = os.path.join(ROOT, "assumptions", f"{slug}.json")
        if os.path.exists(ap):
            data = json.load(open(ap))
            for v in _walk_numbers(data):
                add(v)
        cj = os.path.join(ROOT, "companies", slug, "company.json")
        if os.path.exists(cj):
            for v in _walk_numbers(json.load(open(cj))):
                add(v)
    return known


def _walk_numbers(obj):
    if isinstance(obj, dict):
        for v in obj.values():
            yield from _walk_numbers(v)
    elif isinstance(obj, list):
        for v in obj:
            yield from _walk_numbers(v)
    elif isinstance(obj, (int, float)):
        yield obj


def docx_text(path):
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def xlsx_text(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True)
    out = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for c in row:
                if c is not None:
                    out.append(str(c))
    return "\n".join(out)


def scan(text, known):
    flags = []
    for m in NUM_RE.finditer(text):
        tok = m.group().strip()
        core = normalise(tok)
        if not core or core in ("", "."):
            continue
        try:
            val = float(core)
        except ValueError:
            continue
        # ignore trivially small integers (years, list numbers, single digits, scores)
        if val < 10 and "%" not in tok and "." not in core:
            continue
        if re.match(r"^(19|20)\d\d$", core):  # years
            continue
        if core in known:
            continue
        ctx = text[max(0, m.start() - 14):m.end() + 14]
        if "(e)" in ctx or "~" in ctx or "e)" in ctx:
            continue  # tagged as estimate
        # skip numeric ranges (e.g. 12-18, $70-80, R60-65) and date-adjacent day numbers
        around = text[max(0, m.start() - 2):m.end() + 3]
        if re.search(r"\d\s*[-\u2013]\s*\d", around):
            continue
        if re.search(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)", ctx):
            continue
        flags.append({"token": tok, "context": ctx.replace("\n", " ").strip()})
    # de-dupe by token+context
    seen, uniq = set(), []
    for f in flags:
        k = (f["token"], f["context"])
        if k not in seen:
            seen.add(k); uniq.append(f)
    return uniq


def main():
    ap = argparse.ArgumentParser(description="Lint a deliverable for untagged, unsourced numbers.")
    ap.add_argument("file")
    ap.add_argument("--companies", nargs="+", required=True)
    ap.add_argument("--strict", action="store_true", help="exit non-zero if any flags")
    args = ap.parse_args()
    ext = os.path.splitext(args.file)[1].lower()
    text = docx_text(args.file) if ext == ".docx" else xlsx_text(args.file)
    known = known_numbers(args.companies)
    flags = scan(text, known)
    print(json.dumps({
        "file": args.file,
        "known_actuals_loaded": len(known),
        "flagged_for_review": len(flags),
        "note": "Flags are numbers that are neither a known reported actual nor tagged (e)/~. "
                "Review each: confirm it is sourced, or mark it (e). Heuristic - expect benign hits.",
        "flags": flags[:80],
    }, indent=2))
    if args.strict and flags:
        sys.exit(2)


if __name__ == "__main__":
    main()
