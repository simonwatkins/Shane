#!/usr/bin/env python3
"""House-style analyst-note generator (DOCX) from a structured content spec.

The CHROME of an analyst note (fonts, navy headings, table styling, the scorecard
block, the footer, the disclaimer, the Sources format) is constant; only the
NARRATIVE and the numbers change. This tool renders a per-run content spec (JSON)
into the house-style Word document, so the styling is written and tested once and
the analyst only authors substance.

Built on python-docx (not docx-js) to avoid the truncation and paragraph-border
ordering bugs we hit hand-rolling notes.

Content spec (JSON):
  {
    "title": "...", "subtitle_lines": ["..."], "doc_title": "...", "byline": "...",
    "footer": "...",
    "sections": [
      {"heading": "1. ...", "blocks": [
         {"type": "para", "text": "..."}                  # or "runs": [["txt",{"bold":true}], ...]
         {"type": "h2", "text": "..."},
         {"type": "bullets", "items": ["...", [["a",{"bold":true}],["b",{}]]]},
         {"type": "table", "headers": [...], "rows": [[...]], "col_widths": [...],
          "header_fills": ["NAVY","ACCENT","ACCENT2"]},
         {"type": "scorecard", "ref": "comparisons/shoprite-vs-spar.json"},
         {"type": "callout", "label": "Verdict.", "body": "..."},
         {"type": "sources", "items": [{"label":"...","url":"..."}]}
      ]}
    ]
  }

Usage:
  python3 tools/build_note.py comparisons/shoprite-vs-spar-note.json --out note.docx
"""
import argparse
import json
import os
import sys

from docx import Document
from docx.shared import Pt, Twips
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import report_style as rs

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PALETTE = {"NAVY": rs.NAVY, "ACCENT": rs.ACCENT, "ACCENT2": rs.ACCENT2,
           "GREY": rs.GREY, "RED": rs.RED}


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), rs.FONT); rfonts.set(qn("w:hAnsi"), rs.FONT)
    rPr.append(rfonts)
    col = OxmlElement("w:color"); col.set(qn("w:val"), rs.HYPER); rPr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "21"); rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; r.append(t)
    h.append(r)
    paragraph._p.append(h)


def render_runs(doc, runs, space_after=8):
    rs.docx_rich_para(doc, [(t, o or {}) for t, o in runs], space_after=space_after)


def render_scorecard(doc, ref):
    spec = json.load(open(ref if os.path.isabs(ref) else os.path.join(ROOT, ref)))
    a_code, b_code = spec["companies"]
    headers = ["Resilience dimension", a_code, b_code, "Read (scored for Scenario B, the stress case)"]
    rows, fills = [], {}
    for i, row in enumerate(spec["rows"]):
        rows.append([row["dimension"], str(row["a"]), str(row["b"]),
                     f"{a_code}: {row['why_a']}  /  {b_code}: {row['why_b']}"])
        fills[(i, 1)] = rs.score_fill(row["a"])
        fills[(i, 2)] = rs.score_fill(row["b"])
    # average row
    n = len(spec["rows"])
    avg_a = round(sum(r["a"] for r in spec["rows"]) / n, 1)
    avg_b = round(sum(r["b"] for r in spec["rows"]) / n, 1)
    rows.append(["AVERAGE", str(avg_a), str(avg_b),
                 f"{spec.get('avg_note_a','')}  {spec.get('avg_note_b','')}"])
    fills[(n, 0)] = rs.NAVY
    fills[(n, 1)] = rs.ACCENT
    fills[(n, 2)] = rs.ACCENT2
    t = rs.docx_table(doc, headers, rows, col_widths=[2860, 760, 760, 4980],
                      header_fill=[rs.NAVY, rs.ACCENT, rs.ACCENT2, rs.NAVY],
                      zebra=False, cell_fills=fills, font_size=8.5)
    # white text for the colored score cells + average label
    from docx.shared import RGBColor
    for (i, j), _ in fills.items():
        cell = t.rows[i + 1].cells[j]
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor.from_string("FFFFFF")
                run.bold = True
                if j in (1, 2):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs.docx_para(doc, "", space_after=2)
    return spec


def render_sources(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        rs.docx_run(p, it["label"] + (" " if it.get("url") else ""), size=10.5)
        if it.get("url"):
            add_hyperlink(p, it["url"], it.get("link_text", it["url"]))


def build(spec, out):
    doc = Document()
    rs.apply_house_style(doc)
    sec = doc.sections[0]
    sec.page_width = Twips(12240); sec.page_height = Twips(15840)
    sec.top_margin = Twips(1080); sec.bottom_margin = Twips(1080)
    sec.left_margin = Twips(1440); sec.right_margin = Twips(1440)
    rs.docx_page_number_footer(sec, spec.get("footer", ""))

    # ---- title block
    rs.docx_para(doc, spec["title"], bold=True, size=16, color=rs.NAVY, space_after=2)
    for line in spec.get("subtitle_lines", []):
        rs.docx_para(doc, line, size=10, color="595959", space_after=2)
    if spec.get("doc_title"):
        rs.docx_para(doc, spec["doc_title"], bold=True, size=13, space_after=3)
    if spec.get("byline"):
        rs.docx_para(doc, spec["byline"], italic=True, size=9, color="595959", space_after=10)

    # ---- sections
    for section in spec["sections"]:
        doc.add_heading(section["heading"], level=1)
        for blk in section["blocks"]:
            t = blk["type"]
            if t == "para":
                if "runs" in blk:
                    render_runs(doc, blk["runs"])
                else:
                    rs.docx_para(doc, blk["text"])
            elif t == "h2":
                doc.add_heading(blk["text"], level=2)
            elif t == "bullets":
                for item in blk["items"]:
                    if isinstance(item, str):
                        rs.docx_bullet(doc, item)
                    else:
                        rs.docx_bullet(doc, [(x, o or {}) for x, o in item])
            elif t == "table":
                hf = blk.get("header_fills")
                header_fill = [PALETTE.get(x, rs.NAVY) for x in hf] if hf else rs.NAVY
                rs.docx_table(doc, blk["headers"], blk["rows"],
                              col_widths=blk.get("col_widths"), header_fill=header_fill,
                              zebra=blk.get("zebra", True))
                rs.docx_para(doc, "", space_after=2)
            elif t == "scorecard":
                render_scorecard(doc, blk["ref"])
            elif t == "callout":
                rs.docx_callout(doc, blk.get("label", ""), blk["body"], fill=rs.YELLOW)
                rs.docx_para(doc, "", space_after=2)
            elif t == "sources":
                render_sources(doc, blk["items"])
            else:
                raise SystemExit(f"unknown block type: {t}")

    # ---- standing disclaimer
    rs.docx_para(doc, rs.DISCLAIMER, italic=True, size=8, color="595959", space_before=8)
    doc.save(out)
    _fix_settings_zoom(out)
    print(f"wrote {out}")


def _fix_settings_zoom(path):
    """python-docx emits <w:zoom/> without the required w:percent; patch it so the
    file passes strict OOXML validation."""
    import re
    import shutil
    import tempfile
    import zipfile
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        if "word/settings.xml" not in names:
            return
        data = {n: z.read(n) for n in names}
    xml = data["word/settings.xml"].decode("utf-8")
    if "w:zoom" in xml:
        xml = re.sub(r"<w:zoom\s*/>", '<w:zoom w:percent="100"/>', xml)
        xml = re.sub(r"<w:zoom(?![^>]*w:percent)([^>]*)>", r'<w:zoom w:percent="100"\1>', xml)
        data["word/settings.xml"] = xml.encode("utf-8")
    fd, tmp = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in data.items():
            z.writestr(n, b)
    shutil.move(tmp, path)


def main():
    ap = argparse.ArgumentParser(description="Render a JSE analyst note from a content spec.")
    ap.add_argument("spec")
    ap.add_argument("--out", required=True)
    args = ap.parse_args()
    spec = json.load(open(args.spec if os.path.isabs(args.spec) else os.path.join(ROOT, args.spec)))
    out = args.out if os.path.isabs(args.out) else os.path.join(os.getcwd(), args.out)
    build(spec, out)


if __name__ == "__main__":
    main()
