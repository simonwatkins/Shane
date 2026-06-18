#!/usr/bin/env python3
"""Shared house style for JSE research deliverables (xlsx + docx).

One module, imported by tools/build_model.py and tools/build_note.py, so every
spreadsheet and Word note shares an identical look: Arial, navy headings, the
SHP/SPP-style accent palette, consistent number formats, the standard
"for information only" disclaimer and the Sources format.

Why this exists: the chrome (colours, fonts, table styling, disclaimers) is
constant across runs; only the numbers and narrative change. Keeping the chrome
here means it is written and tested ONCE, instead of being re-derived ad hoc in
every analysis. Mirrors the deterministic discipline of tools/manifest.py.

Deliberately avoids python-docx/openpyxl features that have bitten us before
(e.g. paragraph borders, whose child-element ordering docx engines emit
incorrectly): callout boxes are rendered as single-cell tables instead.
"""

# ----------------------------------------------------------------- palette
FONT = "Arial"

NAVY   = "1F3864"   # primary heading / header bars
ACCENT = "2E5496"   # secondary accent (company A / "SHP blue")
ACCENT2 = "375623"  # tertiary accent (company B / "SPP green")
GREY   = "808080"   # base-case columns / muted
LTGREY = "F2F2F2"   # zebra / label fill
MIDGREY = "D9D9D9"
WHITE  = "FFFFFF"
RED    = "C00000"   # stress scenario / poor score
AMBER  = "BF8F00"   # neutral score
GREEN  = "375623"   # good score
YELLOW = "FFF7E0"   # callout-box fill
HYPER  = "0563C1"

# financial-model text colours (industry convention)
INPUT_BLUE = "0000FF"   # hardcoded inputs / scenario levers
FORMULA_BLACK = "000000" # formulas
LINK_GREEN = "008000"    # cross-sheet links

# ----------------------------------------------------------------- formats
FMT_NUM   = "#,##0"
FMT_NUM1  = "#,##0.0"
FMT_PCT   = "0.0%"
FMT_PCT2  = "0.00%"
FMT_MULT  = '0.00"x"'
FMT_CENT  = '#,##0.0'        # cents value; unit named in the row label
FMT_RATIO = "0.0000"

DISCLAIMER = (
    "For informational purposes only — not investment advice; prepared by internal "
    "research from public disclosures. Figures reported in ZAR unless stated. Reviewed "
    "vs audited status is noted per company. Forward figures and scenario levels are "
    "analyst estimates (e), not company guidance. Share/commodity prices via web search "
    "may be delayed — there is no live JSE feed here. Confirm the fund's dealing-period "
    "status before acting."
)

SACRED_NOTE = (
    "Numbers are sacred: reported actuals are taken verbatim from the cited source; any "
    "figure that is an analyst calculation or projection is marked (e)."
)


def scenario_palette(key):
    """Header fill for a scenario column."""
    k = (key or "").lower()
    if k.startswith("a"):
        return ACCENT
    if k.startswith("b"):
        return RED
    return GREY


def score_fill(score):
    """Colour for a 1-5 resilience score."""
    try:
        s = float(score)
    except (TypeError, ValueError):
        return LTGREY
    if s >= 4:
        return GREEN
    if s >= 3:
        return AMBER
    return RED


# ============================================================ openpyxl helpers
def _xl():
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    return Font, PatternFill, Alignment, Border, Side


def thin_border():
    _, _, _, Border, Side = _xl()
    s = Side(style="thin", color="BFBFBF")
    return Border(left=s, right=s, top=s, bottom=s)


def xl_cell(cell, *, bold=False, size=10, color=FORMULA_BLACK, fill=None,
            align="left", wrap=False, fmt=None, italic=False, border=False):
    """Style one openpyxl cell with house conventions."""
    Font, PatternFill, Alignment, _, _ = _xl()
    cell.font = Font(name=FONT, bold=bold, size=size, color=color, italic=italic)
    cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=wrap)
    if fill:
        cell.fill = PatternFill("solid", start_color=fill)
    if fmt:
        cell.number_format = fmt
    if border:
        cell.border = thin_border()
    return cell


def xl_title_block(ws, title, subtitle, span=9):
    """Navy title bar + accent subtitle across the top of a sheet."""
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=span)
    xl_cell(ws.cell(1, 1, title), bold=True, size=13, color=WHITE, fill=NAVY)
    ws.row_dimensions[1].height = 24
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=span)
    xl_cell(ws.cell(2, 1, subtitle), size=9, color=WHITE, fill=ACCENT, italic=True)
    ws.row_dimensions[2].height = 15


def xl_scenario_header(ws, row, scenario_cols, first_label="R'm unless stated"):
    """Header row: anchor + (label, scenario-key) pairs, colour-coded."""
    xl_cell(ws.cell(row, 2, first_label), bold=True, size=9, color=WHITE,
            fill=NAVY, border=True)
    for j, (label, key) in enumerate(scenario_cols):
        xl_cell(ws.cell(row, 3 + j, label), bold=True, size=9, color=WHITE,
                fill=scenario_palette(key), align="center", wrap=True, border=True)
    ws.row_dimensions[row].height = 26


# ============================================================ python-docx helpers
def _dx():
    from docx.shared import Pt, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    return Pt, RGBColor, Twips, WD_ALIGN_PARAGRAPH


def _rgb(hex6):
    from docx.shared import RGBColor
    return RGBColor.from_string(hex6)


def apply_house_style(doc):
    """Set Normal + heading styles to the house look. Call once on a new Document."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for name, size, color in (("Heading 1", 13, NAVY), ("Heading 2", 11.5, ACCENT)):
        st = doc.styles[name]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = _rgb(color)


def docx_run(p, text, *, bold=False, italic=False, size=10.5, color=None):
    from docx.shared import Pt
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = _rgb(color)
    return r


def docx_para(doc, text="", *, bold=False, italic=False, size=10.5, color=None,
              space_after=6, space_before=0, align=None):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if align == "center":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        docx_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def docx_rich_para(doc, runs, *, space_after=6, space_before=0):
    """runs = list of (text, {opts}) tuples for mixed bold/normal in one paragraph."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    for text, opts in runs:
        docx_run(p, text, **(opts or {}))
    return p


def docx_bullet(doc, text_or_runs):
    from docx.shared import Pt
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if isinstance(text_or_runs, str):
        docx_run(p, text_or_runs)
    else:
        for text, opts in text_or_runs:
            docx_run(p, text, **(opts or {}))
    return p


def _shade_cell(cell, hex6):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex6)
    cell._tc.get_or_add_tcPr().append(sh)


def docx_table(doc, headers, rows, *, col_widths=None, header_fill=NAVY,
               zebra=True, cell_fills=None, font_size=9.0):
    """Styled table. headers=list[str]; rows=list[list[str]].

    cell_fills: optional dict {(row_idx, col_idx): hex} for per-cell shading
    (row_idx is 0-based over data rows). Header cells: white bold on header_fill.
    """
    from docx.shared import Pt, Twips
    from docx.enum.table import WD_TABLE_ALIGNMENT
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = ""
        p = hdr[j].paragraphs[0]
        docx_run(p, h, bold=True, size=font_size, color=WHITE)
        _shade_cell(hdr[j], header_fill if not isinstance(header_fill, (list, tuple))
                    else header_fill[j])
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            docx_run(p, str(val), size=font_size, bold=(j == 0))
            fill = None
            if cell_fills and (i, j) in cell_fills:
                fill = cell_fills[(i, j)]
            elif j == 0:
                fill = LTGREY
            elif zebra and i % 2 == 1:
                fill = "FAFAFA"
            if fill:
                _shade_cell(cells[j], fill)
    if col_widths:
        for j, w in enumerate(col_widths):
            for r in t.rows:
                r.cells[j].width = Twips(w)
    return t


def docx_callout(doc, label, body, fill=YELLOW):
    """A shaded single-cell callout box (used instead of paragraph borders)."""
    from docx.shared import Twips, Pt
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell.width = Twips(9360)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if label:
        docx_run(p, label + "  ", bold=True, size=10.5)
    docx_run(p, body, size=10.5)
    _shade_cell(cell, fill)
    return t


def docx_page_number_footer(section, prefix=""):
    """Add a centred footer with 'prefix Page N' using a PAGE field."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if prefix:
        docx_run(p, prefix + "  Page ", size=7.5, color=GREY)
    else:
        docx_run(p, "Page ", size=7.5, color=GREY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "15"); rpr.append(sz)
    r.append(rpr)
    fld.append(r)
    p._p.append(fld)
    return footer
